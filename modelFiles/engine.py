"""
Clean, importable module extracted from the Colab notebook jrv2 (1).py.
Contains all model classes and utilities needed for job recommendation inference.

Usage:
    from engine import ConvDeepFMJobRecommender
    engine = ConvDeepFMJobRecommender(model_dir="./", groq_api_key="...")
    result = engine.recommend_from_resume(resume_path, resume_type, location, ...)
"""

import torch
import torch.nn as nn
import numpy as np
import pandas as pd
import json
import re
import os
import sys
import pickle
import unicodedata
from collections import Counter
from sklearn.preprocessing import LabelEncoder
from sklearn.feature_extraction.text import TfidfVectorizer
from tqdm import tqdm
from datetime import datetime, timedelta

import pdfplumber
import docx
# jobspy is imported lazily inside JobProcessor.scrape_jobs() to avoid
# a native crash from tls_client on some Windows configurations.
from sentence_transformers import SentenceTransformer
from groq import Groq

# ── Config ──────────────────────────────────────────────────────────────────
GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")
GROQ_MODEL = "llama-3.3-70b-versatile"

device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')

# ============================================================================
# SKILLS DATABASE
# ============================================================================
SKILLS_DATABASE = [
    # Programming Languages
    'Python', 'Java', 'JavaScript', 'TypeScript', 'C++', 'C#', 'C', 'PHP',
    'Ruby', 'Swift', 'Kotlin', 'Go', 'Rust', 'Scala', 'R', 'MATLAB',
    'Perl', 'Bash', 'Shell', 'PowerShell', 'Lua', 'Julia', 'Haskell',
    'Elixir', 'Erlang', 'Clojure', 'F#', 'Dart', 'Groovy', 'COBOL',
    'Fortran', 'Assembly', 'Objective-C', 'VBA', 'SAS', 'Stata',
    # Web Frontend
    'HTML', 'CSS', 'React', 'Angular', 'Vue.js', 'jQuery', 'Bootstrap',
    'Tailwind CSS', 'SASS', 'LESS', 'Redux', 'Next.js', 'Nuxt.js',
    'Svelte', 'Gatsby', 'Webpack', 'Vite', 'Babel', 'Three.js',
    'WebAssembly', 'PWA', 'Storybook', 'Figma', 'Adobe XD',
    # Web Backend
    'Node.js', 'Express.js', 'Django', 'Flask', 'FastAPI', 'Spring Boot',
    'ASP.NET', 'Ruby on Rails', 'Laravel', 'NestJS', 'Gin', 'Echo',
    'Phoenix', 'Fiber', 'Actix', 'Hapi.js', 'Koa.js', 'Strapi',
    'GraphQL', 'REST API', 'gRPC', 'WebSocket', 'Microservices',
    # Databases
    'MySQL', 'PostgreSQL', 'MongoDB', 'Redis', 'Oracle', 'SQLite',
    'Cassandra', 'DynamoDB', 'Elasticsearch', 'Neo4j', 'MariaDB',
    'CouchDB', 'InfluxDB', 'TimescaleDB', 'Firestore', 'Supabase',
    'SQL', 'NoSQL', 'PL/SQL', 'T-SQL',
    # Data Science & ML
    'TensorFlow', 'PyTorch', 'Keras', 'Scikit-learn', 'Pandas', 'NumPy',
    'Matplotlib', 'Seaborn', 'OpenCV', 'NLTK', 'SpaCy', 'Transformers',
    'Machine Learning', 'Deep Learning', 'Neural Networks', 'NLP',
    'Computer Vision', 'Data Analysis', 'Statistical Analysis',
    'Reinforcement Learning', 'Generative AI', 'LLM', 'RAG',
    'Time Series Analysis', 'Anomaly Detection', 'Recommendation Systems',
    'A/B Testing', 'Hypothesis Testing', 'Regression Analysis',
    'Classification', 'Clustering', 'Dimensionality Reduction',
    'Feature Engineering', 'Model Deployment', 'Model Evaluation',
    'XGBoost', 'LightGBM', 'CatBoost', 'Random Forest',
    # Big Data & Data Engineering
    'Spark', 'Hadoop', 'Kafka', 'Airflow', 'Databricks', 'Hive',
    'Flink', 'Presto', 'Trino', 'dbt', 'Fivetran', 'Stitch',
    'Apache Beam', 'Apache NiFi', 'Delta Lake', 'Iceberg',
    'ETL', 'ELT', 'Data Warehousing', 'Data Pipeline', 'Data Modeling',
    'Redshift', 'BigQuery', 'Snowflake', 'Looker', 'Dataflow',
    # Cloud
    'AWS', 'Azure', 'GCP', 'Heroku', 'DigitalOcean', 'Firebase',
    'Cloudflare', 'Vercel', 'Netlify',
    'AWS Lambda', 'AWS EC2', 'AWS S3', 'AWS RDS', 'AWS EKS',
    'Azure DevOps', 'Azure Functions', 'Google Cloud Run',
    # DevOps
    'Docker', 'Kubernetes', 'Jenkins', 'Git', 'GitHub', 'GitLab',
    'CI/CD', 'Terraform', 'Ansible', 'Linux', 'Unix', 'Nginx',
    'Prometheus', 'Grafana', 'ELK Stack', 'Datadog', 'New Relic',
    'SonarQube', 'ArgoCD', 'Helm', 'Istio', 'Consul', 'Vault',
    'GitHub Actions', 'CircleCI', 'Travis CI', 'Bitbucket',
    # Mobile
    'Android', 'iOS', 'React Native', 'Flutter', 'Xamarin', 'Ionic',
    'Swift UI', 'Jetpack Compose', 'Expo', 'Capacitor',
    # MLOps & AI
    'MLflow', 'Kubeflow', 'SageMaker', 'Vertex AI', 'Ray',
    'HuggingFace', 'LangChain', 'LlamaIndex', 'Weights & Biases',
    'DVC', 'BentoML', 'TorchServe', 'TFServing',
    'OpenAI API', 'Anthropic API', 'Cohere', 'Pinecone', 'Weaviate',
    'ChromaDB', 'FAISS', 'Qdrant',
    # BI & Viz
    'Tableau', 'Power BI', 'Metabase', 'Superset', 'D3.js', 'Plotly',
    'Dash', 'Kibana', 'Google Data Studio', 'Qlik',
    # Security
    'Cybersecurity', 'Penetration Testing', 'Ethical Hacking',
    'Network Security', 'Firewall', 'VPN', 'SIEM', 'SOC',
    'OWASP', 'SSL/TLS', 'IAM', 'OAuth', 'JWT', 'Encryption',
    'Wireshark', 'Metasploit', 'Burp Suite', 'Nmap', 'Kali Linux',
    'ISO 27001', 'NIST', 'SOC 2', 'Zero Trust',
    # Testing
    'Selenium', 'Cypress', 'Jest', 'Mocha', 'Pytest', 'JUnit',
    'Playwright', 'TestNG', 'Appium', 'Postman',
    'Unit Testing', 'Integration Testing', 'Load Testing', 'QA',
    # Project Tools
    'Agile', 'Scrum', 'Kanban', 'JIRA', 'Confluence', 'Notion',
    'Slack', 'Trello', 'Asana', 'Monday.com', 'ClickUp',
    # Design
    'Figma', 'Adobe XD', 'Sketch', 'InVision',
    'UI Design', 'UX Design', 'Wireframing', 'Prototyping',
    'User Research', 'Usability Testing', 'Design Systems',
    'Adobe Photoshop', 'Adobe Illustrator', 'Canva',
    # Blockchain
    'Blockchain', 'Solidity', 'Ethereum', 'Web3.js', 'Hardhat',
    'Smart Contracts', 'DeFi', 'NFT', 'IPFS', 'Polygon', 'Solana',
    # Game Dev
    'Unity', 'Unreal Engine', 'Godot', 'Game Design', 'Level Design',
    '3D Modeling', 'Blender',
    # Embedded
    'Embedded Systems', 'Arduino', 'Raspberry Pi', 'FPGA',
    'VHDL', 'Verilog', 'IoT', 'RTOS', 'Microcontrollers',
    'PLC', 'SCADA', 'Industrial Automation',
    # Finance
    'Excel', 'Microsoft Excel', 'Advanced Excel', 'Tally', 'Tally Prime',
    'QuickBooks', 'Zoho Books', 'SAP', 'SAP FICO', 'Oracle Financials',
    'VLOOKUP', 'Pivot Tables', 'Power Query', 'Macros', 'VBA',
    'Financial Analysis', 'Financial Modeling', 'Financial Reporting',
    'Accounting', 'Bookkeeping', 'Auditing', 'Internal Audit',
    'Taxation', 'GST', 'Income Tax', 'TDS',
    'Cost Accounting', 'Management Accounting', 'Budgeting', 'Forecasting',
    'Investment Analysis', 'Equity Research', 'Portfolio Management',
    'Risk Management', 'Credit Analysis', 'Derivatives',
    'Fundamental Analysis', 'Technical Analysis', 'Valuation',
    'Investment Banking', 'Corporate Finance', 'Treasury Management',
    'IFRS', 'GAAP', 'CFA', 'CA', 'CPA', 'ACCA',
    'Bloomberg Terminal', 'Reuters Eikon',
    # Marketing & Sales
    'Digital Marketing', 'SEO', 'SEM', 'Google Ads', 'Facebook Ads',
    'Content Marketing', 'Email Marketing', 'Social Media Marketing',
    'Brand Management', 'CRM', 'Salesforce', 'HubSpot',
    'Lead Generation', 'Sales', 'B2B Sales', 'B2C Sales',
    'Business Development', 'Copywriting', 'Content Writing',
    'Google Analytics', 'Marketing Automation', 'Mailchimp', 'Marketo',
    'E-commerce', 'Shopify', 'WooCommerce', 'Amazon FBA',
    # HR
    'Human Resources', 'HR', 'Recruitment', 'Talent Acquisition',
    'Payroll', 'HRMS', 'HRIS', 'Workday', 'BambooHR', 'Darwinbox',
    'Performance Management', 'KPI', 'OKR', 'Employee Engagement',
    'Training & Development', 'L&D', 'Compensation & Benefits',
    'HR Analytics', 'Organizational Development', 'Labor Law',
    # Operations
    'Operations Management', 'Supply Chain', 'Logistics', 'Procurement',
    'Vendor Management', 'Inventory Management', 'ERP',
    'Warehouse Management', 'Quality Control', 'Quality Assurance',
    'Six Sigma', 'Lean', 'ISO 9001', 'SAP SCM',
    # PM
    'Project Management', 'Product Management', 'Program Management',
    'PMP', 'Prince2', 'PMI', 'Stakeholder Management',
    'Business Analysis', 'Requirements Gathering', 'Product Roadmap',
    # Healthcare
    'Clinical Research', 'Clinical Trials', 'GCP', 'Pharmacovigilance',
    'Regulatory Affairs', 'FDA', 'Medical Writing', 'Biostatistics',
    'Healthcare Management', 'EMR', 'EHR', 'HL7', 'FHIR',
    'Medical Coding', 'ICD-10', 'MBBS', 'BPharm', 'MPharm',
    # Education
    'Teaching', 'Curriculum Development', 'Instructional Design',
    'E-Learning', 'LMS', 'Moodle', 'Academic Research',
    # Legal
    'Legal Research', 'Contract Drafting', 'Corporate Law',
    'Intellectual Property', 'GDPR', 'Data Privacy', 'Regulatory Compliance',
    'Arbitration', 'Mediation', 'LLB', 'LLM',
    # Engineering (non-software)
    'AutoCAD', 'SolidWorks', 'CATIA', 'ANSYS', 'MATLAB', 'ETABS',
    'Mechanical Engineering', 'Civil Engineering', 'FEA', 'CFD',
    'BIM', 'Revit', 'HVAC', 'Electrical Engineering', 'PLC',
    # Creative
    'Graphic Design', 'Video Production', 'Photography', 'After Effects',
    'Adobe Premiere Pro', 'DaVinci Resolve', 'Journalism', 'Scriptwriting',
    # Other
    'Hotel Management', 'Event Management', 'Travel Management',
    'Personal Training', 'Yoga', 'Nutrition', 'Childcare', 'Fashion Design',
]

# ============================================================================
# DOMAIN KEYWORDS
# ============================================================================
DOMAIN_KEYWORDS = {
    'Data Science':       ['machine learning', 'deep learning', 'nlp', 'tensorflow', 'pytorch',
                           'pandas', 'numpy', 'scikit-learn', 'data analysis', 'computer vision',
                           'statistical analysis', 'xgboost', 'lightgbm', 'generative ai', 'llm'],
    'Web Development':    ['react', 'angular', 'vue', 'node.js', 'django', 'flask',
                           'html', 'css', 'javascript', 'typescript', 'next.js', 'svelte'],
    'DevOps':             ['docker', 'kubernetes', 'jenkins', 'ci/cd', 'terraform',
                           'ansible', 'linux', 'aws', 'azure', 'gcp', 'prometheus', 'helm'],
    'Mobile Development': ['android', 'ios', 'flutter', 'react native', 'swift', 'kotlin',
                           'jetpack compose', 'swift ui', 'expo'],
    'Backend':            ['spring boot', 'fastapi', 'microservices', 'rest api', 'grpc',
                           'postgresql', 'mongodb', 'redis', 'kafka', 'rabbitmq'],
    'Data Engineering':   ['spark', 'hadoop', 'airflow', 'dbt', 'databricks', 'kafka',
                           'etl', 'elt', 'data pipeline', 'data warehousing', 'snowflake',
                           'bigquery', 'redshift', 'delta lake'],
    'Cybersecurity':      ['cybersecurity', 'penetration testing', 'ethical hacking',
                           'network security', 'siem', 'soc', 'owasp', 'vulnerability assessment',
                           'incident response', 'kali linux', 'burp suite'],
    'UI/UX Design':       ['figma', 'adobe xd', 'ui design', 'ux design', 'wireframing',
                           'prototyping', 'user research', 'usability testing', 'sketch'],
    'Blockchain':         ['blockchain', 'solidity', 'ethereum', 'web3', 'smart contracts',
                           'defi', 'nft', 'polygon', 'solana'],
    'Game Development':   ['unity', 'unreal engine', 'game design', 'godot',
                           'level design', '3d modeling', 'blender'],
    'Embedded Systems':   ['embedded systems', 'arduino', 'raspberry pi', 'fpga', 'iot',
                           'rtos', 'microcontrollers', 'plc', 'scada', 'vhdl', 'verilog'],
    'Finance':            ['financial analysis', 'financial modeling', 'investment analysis',
                           'equity research', 'portfolio management', 'trading',
                           'tally', 'budgeting', 'forecasting', 'mis reporting', 'cfa',
                           'investment banking', 'corporate finance', 'valuation', 'bloomberg'],
    'Accounting':         ['accounting', 'bookkeeping', 'auditing', 'taxation', 'gst',
                           'tally', 'quickbooks', 'sap fico', 'cost accounting',
                           'financial statements', 'ifrs', 'gaap', 'ca', 'cpa', 'acca'],
    'Marketing':          ['digital marketing', 'seo', 'sem', 'google ads', 'content marketing',
                           'email marketing', 'social media marketing', 'brand management',
                           'performance marketing', 'google analytics', 'marketing automation'],
    'Sales':              ['sales', 'lead generation', 'crm', 'salesforce', 'hubspot',
                           'b2b sales', 'b2c sales', 'account management', 'business development'],
    'HR':                 ['human resources', 'recruitment', 'talent acquisition', 'payroll',
                           'hrms', 'workday', 'performance management', 'l&d', 'hr analytics'],
    'Operations':         ['operations management', 'supply chain', 'logistics', 'procurement',
                           'vendor management', 'inventory management', 'erp', 'six sigma',
                           'lean', 'iso 9001', 'demand planning'],
    'Project Management': ['project management', 'pmp', 'agile', 'scrum', 'kanban',
                           'stakeholder management', 'risk management', 'program management',
                           'safe', 'prince2'],
    'Product Management': ['product management', 'product roadmap', 'product strategy',
                           'backlog grooming', 'sprint planning', 'okr', 'go-to-market'],
    'Consulting':         ['management consulting', 'strategy consulting', 'business strategy',
                           'digital transformation', 'process improvement', 'swot analysis',
                           'competitive intelligence', 'business case'],
    'Business Analysis':  ['business analysis', 'requirements gathering', 'brd', 'frd',
                           'market research', 'data analysis', 'process mapping'],
    'Legal':              ['legal research', 'contract drafting', 'corporate law', 'litigation',
                           'intellectual property', 'gdpr', 'compliance', 'arbitration', 'llb'],
    'Healthcare':         ['clinical research', 'pharmacovigilance', 'regulatory affairs',
                           'medical writing', 'biostatistics', 'ehr', 'fhir', 'medical coding'],
    'Education':          ['teaching', 'curriculum development', 'instructional design',
                           'e-learning', 'academic research', 'lms', 'thesis writing'],
    'Engineering':        ['autocad', 'solidworks', 'catia', 'ansys', 'mechanical engineering',
                           'civil engineering', 'fea', 'cfd', 'bim', 'revit', 'hvac'],
    'Creative':           ['graphic design', 'video production', 'photography', 'figma',
                           'adobe photoshop', 'illustrator', 'motion graphics', 'after effects',
                           'content writing', 'copywriting', 'journalism'],
    'Hospitality':        ['hotel management', 'front office', 'revenue management',
                           'guest relations', 'event planning', 'travel management', 'amadeus'],
    'Fitness & Wellness': ['personal training', 'fitness training', 'group fitness',
                           'nutrition', 'diet planning', 'weight loss', 'yoga', 'pilates',
                           'kinesiology', 'cpr', 'ace certification', 'nasm', 'wellness coaching'],
    'Media & Film':       ['video production', 'video editing', 'adobe premiere pro',
                           'davinci resolve', 'after effects', 'cinematography',
                           'scriptwriting', 'audio editing', 'content creation', 'podcast production'],
    'Fashion':            ['fashion design', 'pattern making', 'fabric selection',
                           'sketching', 'merchandising', 'garment construction', 'visual merchandising'],
    'Nutrition':          ['nutritional counseling', 'dietary assessment', 'meal planning',
                           'food science', 'clinical nutrition', 'dietetics', 'registered dietitian'],
    'Account Management': ['account management', 'key account management', 'global accounts',
                           'client relationship management', 'customer success', 'upselling'],
}

VALID_DOMAINS = set(DOMAIN_KEYWORDS.keys())


# ============================================================================
# Groq LLM Client
# ============================================================================
class GroqLLMClient:
    def __init__(self, api_key: str, model: str = GROQ_MODEL):
        import time as _time
        self._time = _time
        # Strip whitespace/newlines — a line break in the key causes
        # httpcore.LocalProtocolError: Illegal header value
        api_key = "".join(api_key.split())
        self.client = Groq(
            api_key=api_key,
            timeout=120.0,      # 120s overall (default connect=5s is too tight)
            max_retries=3,      # retry transient connection failures
        )
        self.model = model
        self._available = True

    def _call(self, system_prompt: str, user_prompt: str,
              max_tokens: int = 1024, temperature: float = 0.1) -> str | None:
        if not self._available:
            print("[GROQ] _call skipped — client marked unavailable")
            return None
        last_err = None
        for attempt in range(1, 4):
            try:
                t0 = self._time.time()
                preview = user_prompt[:120].replace('\n', ' ')
                print(f"[GROQ] _call attempt {attempt}/3 — model={self.model}, "
                      f"max_tokens={max_tokens}, prompt_len={len(user_prompt)} chars ...")
                print(f"[GROQ]   prompt preview: {preview!r}...")
                resp = self.client.chat.completions.create(
                    model=self.model,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt},   # ← FULL prompt (was wrongly [:80])
                    ],
                    max_tokens=max_tokens,
                    temperature=temperature,
                )
                elapsed = self._time.time() - t0
                content = resp.choices[0].message.content.strip()
                print(f"[GROQ] _call SUCCESS in {elapsed:.1f}s — response length: {len(content)} chars")
                return content
            except RuntimeError as e:
                print(f"[GROQ] RuntimeError (event loop conflict): {e}")
                import traceback; traceback.print_exc()
                return None
            except Exception as e:
                elapsed = self._time.time() - t0
                last_err = e
                print(f"[GROQ] ERROR attempt {attempt}/3 after {elapsed:.1f}s — {type(e).__name__}: {e}")
                import traceback; traceback.print_exc()
                if attempt < 3:
                    wait = 2 * attempt
                    print(f"[GROQ] Retrying in {wait}s ...")
                    self._time.sleep(wait)
                    continue
        print(f"[GROQ] FAILED after 3 attempts. Last error: {type(last_err).__name__}: {last_err}")
        return None

    @staticmethod
    def _parse_json(text: str) -> dict | list | None:
        """
        Robustly extract JSON from model output.
        Handles: markdown fences, trailing commas, truncated arrays, wrapped objects.
        """
        if not text:
            return None
        # Strip markdown fences
        text = re.sub(r'^```(?:json)?\s*', '', text.strip(), flags=re.I)
        text = re.sub(r'\s*```$', '', text)
        text = text.strip()

        # Attempt 1: clean parse
        try:
            return json.loads(text)
        except json.JSONDecodeError:
            pass

        # Attempt 2: strip trailing commas before } or ] (common LLM mistake)
        cleaned = re.sub(r',\s*([}\]])', r'\1', text)
        try:
            return json.loads(cleaned)
        except json.JSONDecodeError:
            pass

        # Attempt 3: extract first {...} or [...] block
        m = re.search(r'(\{.*\}|\[.*\])', cleaned, re.S)
        if m:
            try:
                return json.loads(m.group(1))
            except json.JSONDecodeError:
                pass

        # Attempt 4: recover truncated JSON arrays by closing them
        if cleaned.lstrip().startswith('['):
            # Find last complete object in the array
            last_close = cleaned.rfind('}')
            if last_close != -1:
                candidate = cleaned[:last_close + 1] + ']'
                try:
                    return json.loads(candidate)
                except json.JSONDecodeError:
                    pass

        return None

    @staticmethod
    def _keyword_score_domains(skills: list[str], top_n: int = 3) -> list[tuple[str, int]]:
        sl = [s.lower() for s in skills]
        scores = {
            d: sum(1 for kw in kws if any(kw in s for s in sl))
            for d, kws in DOMAIN_KEYWORDS.items()
        }
        ranked = sorted(scores.items(), key=lambda x: x[1], reverse=True)
        top = [(d, sc) for d, sc in ranked if sc > 0][:top_n]
        return top if top else [('Software Engineer', 0)]

    def validate_domains(self, llm_domains: list[str], skills: list[str]) -> list[tuple[str, int]]:
        sl = [s.lower() for s in skills]
        accepted = []
        dropped = []

        for domain in llm_domains:
            domain = domain.strip()
            if not domain:
                continue
            if domain in VALID_DOMAINS:
                canonical = domain
            else:
                canonical = next(
                    (d for d in VALID_DOMAINS if d.lower() == domain.lower()), None
                )
            if canonical is None:
                dropped.append(domain)
                continue
            kws = DOMAIN_KEYWORDS[canonical]
            score = sum(1 for kw in kws if any(kw in s for s in sl))
            accepted.append((canonical, score))

        if accepted:
            seen = set()
            unique = []
            for d, sc in accepted:
                if d not in seen:
                    seen.add(d)
                    unique.append((d, sc))
            return unique[:3]

        return self._keyword_score_domains(skills)

    def parse_resume_llm(self, resume_text: str) -> dict | None:
        print(f"[GROQ] parse_resume_llm called — resume text length: {len(resume_text)} chars")
        system = (
            "You are a world-class resume-parsing and career-intelligence engine with deep expertise "
            "across every industry — technology, finance, healthcare, law, marketing, engineering, and more. "
            "\n\nYour task: read the full resume carefully and extract a rich, accurate profile. "
            "Think step-by-step before answering: read the entire text, identify work history dates "
            "precisely, enumerate every skill and tool mentioned, infer the seniority level from "
            "titles and years of experience, and select the most precise job domain(s)."
            "\n\nOutput rules (STRICT):"
            "\n• Return ONLY a single valid JSON object."
            "\n• No markdown fences (no ```json), no commentary, no preamble."
            "\n• All string values must be properly JSON-escaped."
            "\n• If a field cannot be determined, use its specified default."
        )
        valid_domains_str = ", ".join(sorted(VALID_DOMAINS))
        # Send as much of the resume as the context window allows (llama-3.3-70b has 128k ctx)
        resume_snippet = resume_text[:14000]
        user = f"""Carefully read the resume below and return a JSON object with EXACTLY these keys.

=== OUTPUT FORMAT ===
{{
  "skills": [
    "<every technical skill, tool, language, framework, library, platform, methodology, "
    "certification, domain-specific knowledge, and relevant soft skill found in the resume. "
    "Be exhaustive — list everything explicitly mentioned or strongly implied. "
    "Use canonical names: e.g. 'PostgreSQL' not 'postgres', 'Scikit-learn' not 'sklearn'."
  ],
  "experience_years": <float — sum of ALL non-overlapping professional work periods in years,
                       calculated precisely from date ranges. Count internships if listed.
                       Return 0.0 for freshers / students with no work experience.
                       Example: Jan 2021–Jun 2022 (1.5yr) + Aug 2022–present (2.7yr) = 4.2>,
  "education_level": "<EXACT highest qualification from the resume, e.g.:
                       'B.Tech in Computer Science', 'MBA in Finance', 'M.Sc in Data Science',
                       'CA (Chartered Accountant)', 'MBBS', 'LLB', 'PhD in Machine Learning'>",
  "domains": [
    "<1 to 3 domains that BEST match this candidate's overall profile and career trajectory.
     MUST be chosen ONLY from this list: {valid_domains_str}.
     Order them from best match to weakest. Never invent new domain names.>"
  ],
  "job_title": "<the single most precise, realistic job title for this candidate based on their
                 most recent role, skills, and experience level. Be specific:
                 e.g. 'Senior Machine Learning Engineer', 'Chartered Accountant', 'Full-Stack Developer',
                 'Digital Marketing Manager' — NOT generic terms like 'Engineer' or 'Executive'>",
  "seniority": "<one of: Fresher | Junior | Mid-level | Senior | Lead | Manager | Director | Executive>",
  "location": "<candidate's current city/state/country from the resume. Look for addresses, phone area codes, or stated locations. e.g. 'Bangalore, India', 'San Francisco, CA'. Return empty string if not found.>"
}}

=== RULES ===
1. skills — include EVERYTHING: programming languages, frameworks, databases, cloud platforms,
   DevOps tools, BI tools, methodologies (Agile/Scrum), certifications (AWS Certified, PMP, CFA, CA),
   domain knowledge (GST, IFRS, HL7), soft skills if explicitly mentioned (Leadership, Communication).
2. experience_years — compute from actual date ranges in the resume. DO NOT just read a stated
   summary like "5 years experience" — verify it from the work history dates.
3. education_level — use the exact degree name as written in the resume.
4. domains — think holistically about the candidate's CAREER, not just one resume section.
5. job_title — match the candidate's actual seniority and specialisation precisely.

=== RESUME TEXT ===
\"\"\"
{resume_snippet}
\"\"\"
"""
        raw = self._call(system, user, max_tokens=2000, temperature=0.05)
        if not raw:
            print("[GROQ] parse_resume_llm — _call returned None (Groq failed)")
            return None
        parsed = self._parse_json(raw)
        if not isinstance(parsed, dict):
            print(f"[GROQ] parse_resume_llm — JSON parse failed, got: {type(parsed)}")
            return None
        print(f"[GROQ] parse_resume_llm — SUCCESS, extracted {len(parsed.get('skills', []))} skills")
        result = {
            'skills': parsed.get('skills', []) or [],
            'experience_years': float(parsed.get('experience_years', 0) or 0),
            'education_level': str(parsed.get('education_level', 'B.Tech in Computer Science')),
            'domains_raw': parsed.get('domains', []) or [],
            'job_title': str(parsed.get('job_title', '')),
            'seniority': str(parsed.get('seniority', '')),
            'location': str(parsed.get('location', '')),
        }
        if not isinstance(result['skills'], list):
            result['skills'] = []
        if not isinstance(result['domains_raw'], list):
            result['domains_raw'] = []
        return result

    def extract_job_skills_llm(self, job_descriptions: list[dict]) -> dict[str, list[str]]:
        print(f"[GROQ] extract_job_skills_llm called — {len(job_descriptions)} jobs to enrich")
        if not self._available:
            print("[GROQ] extract_job_skills_llm skipped — client unavailable")
            return {}
        system = (
            "You are an expert talent acquisition specialist and job description analyst with "
            "deep knowledge across all industries — technology, finance, healthcare, legal, "
            "marketing, engineering, operations, and more."
            "\n\nYour task: for each job description provided, extract a COMPREHENSIVE list of "
            "skills, qualifications, tools, and knowledge areas that a successful candidate needs."
            "\n\nExtraction rules:"
            "\n• Include: programming languages, frameworks, libraries, cloud platforms, databases, "
            "DevOps tools, BI/analytics tools, industry-specific software, certifications, "
            "methodologies (Agile, Six Sigma, IFRS), domain knowledge (GST, FHIR, GAAP), "
            "and any explicitly required soft skills (Leadership, Communication, etc.)."
            "\n• Use canonical skill names: e.g. 'PostgreSQL' not 'Postgres', 'Scikit-learn' not 'sklearn'."
            "\n• Do NOT include vague non-skills like 'passion', 'attitude', 'team player' "
            "unless they are explicitly stated as requirements."
            "\n• Return ONLY a valid JSON object — no markdown fences, no commentary, no preamble."
        )
        results = {}
        batch_size = 4  # smaller batches = larger description window per job = better extraction
        items = list(job_descriptions)
        total_batches = (len(items) + batch_size - 1) // batch_size
        for batch_start in range(0, len(items), batch_size):
            batch_num = batch_start // batch_size + 1
            batch = items[batch_start: batch_start + batch_size]
            print(f"[GROQ] extract_job_skills_llm — batch {batch_num}/{total_batches} ({len(batch)} jobs)")
            jobs_block = ""
            for entry in batch:
                jid = entry['job_id']
                desc = entry['description'][:2000]  # was 800 — now 2000 for much richer extraction
                jobs_block += (
                    f'\n### JOB_ID: {jid}\n'
                    f'DESCRIPTION:\n"""\n{desc}\n"""\n'
                )
            user = f"""For each job below, extract a COMPREHENSIVE list of required skills and qualifications.

Return a JSON object where:
- Each KEY is the exact JOB_ID string (e.g. "job_0")
- Each VALUE is a flat list of skill strings

Important:
• Be exhaustive — capture every technical skill, tool, certification, methodology, and domain knowledge mentioned.
• Skills should be specific and canonical (e.g. "React" not "React framework", "AWS" not "Amazon cloud").
• For non-tech roles include domain skills too (e.g. ["GST", "Tally Prime", "Financial Reporting", "IFRS"] for finance).
• Minimum 5 skills per job if the description has enough content. Aim for 10–20 skills per job.

Example output format:
{{
  "job_0": ["Python", "TensorFlow", "MLflow", "AWS SageMaker", "SQL", "Docker", "Machine Learning", "Data Analysis"],
  "job_1": ["Financial Analysis", "Excel", "CFA", "Bloomberg Terminal", "Equity Research", "Valuation", "Financial Modeling"],
  "job_2": ["Patient Care", "MBBS", "Laparoscopy", "Clinical Research", "EHR", "ICD-10", "Medical Writing"]
}}

{jobs_block}
"""
            raw = self._call(system, user, max_tokens=2000, temperature=0.0)
            if not raw:
                continue
            parsed = self._parse_json(raw)
            if isinstance(parsed, dict):
                for jid, skills in parsed.items():
                    if isinstance(skills, list):
                        results[jid] = [str(s).strip() for s in skills if s]
        return results

    def score_jobs_with_llm(self, user_profile: dict, jobs: list[dict],
                             batch_size: int = 8) -> dict[str, float]:
        """
        Use Groq to directly rate how well each job fits the candidate.

        This is the 4th and most powerful scoring signal — Groq understands
        nuance that keyword overlap and BERT embeddings cannot:
        - Implicit skill matches (e.g. "built NLP pipelines" → NLP Engineer role)
        - Seniority alignment (Senior vs Junior titles)
        - Industry context (e.g. fintech background → finance-adjacent tech roles)
        - Career trajectory logic (e.g. BA → PM progression)

        Strategy:
        - Build a concise but rich candidate summary once.
        - Send jobs in small batches (to maximise description length per job).
        - Groq returns a JSON array: [{job_id, score, reason}]
        - Scores are normalised to [0.0, 1.0].
        - Falls back gracefully to empty dict on failure (caller uses 0.0).

        Args:
            user_profile: parsed resume dict (skills, experience_years, etc.)
            jobs: list of processed job dicts (job_id, title, company, full_description, skills)
            batch_size: jobs per Groq call (default 8, trades throughput vs token budget)

        Returns:
            dict of {job_id: float score 0.0–1.0}
        """
        if not self._available:
            print("[GROQ] score_jobs_with_llm skipped — client unavailable")
            return {}

        llm_scores: dict[str, float] = {}

        # Build a rich but compact candidate snapshot (sent with every batch)
        skills_str = ", ".join(user_profile.get('skills', [])[:30])
        resume_snippet = user_profile.get('resume_text', '')[:1200].strip()
        seniority = user_profile.get('seniority', '')
        candidate_summary = (
            f"CANDIDATE PROFILE\n"
            f"Job Title       : {user_profile.get('job_title', 'Not specified')}\n"
            f"Seniority       : {seniority or 'Not specified'}\n"
            f"Primary Domain  : {user_profile.get('domain', 'Not specified')}\n"
            f"Experience      : {user_profile.get('experience_years', 0)} years\n"
            f"Education       : {user_profile.get('education_level', 'Not specified')}\n"
            f"Key Skills      : {skills_str or 'None listed'}\n"
            f"Resume Excerpt  :\n{resume_snippet}"
        )

        system = (
            "You are a senior recruitment specialist with 20 years of experience across "
            "technology, finance, healthcare, law, marketing, and all other industries. "
            "Your task: given a candidate profile and a list of job postings, assign each job "
            "a relevance score from 0.0 to 1.0 indicating how well the candidate fits that role."
            "\n\nScoring guide:"
            "\n  1.0  — Perfect fit: skills, experience, seniority and domain all align tightly."
            "\n  0.8  — Strong fit: minor gaps (e.g. 1 tool missing, slightly over/under-experienced)."
            "\n  0.6  — Decent fit: transferable skills, adjacent domain, or slight seniority mismatch."
            "\n  0.4  — Weak fit: some relevance but meaningful skill or domain gaps."
            "\n  0.2  — Poor fit: barely relevant, major gaps."
            "\n  0.0  — No fit: completely unrelated role."
            "\n\nBe nuanced. Use context from the resume excerpt, not just keyword matching. "
            "Consider career trajectory, domain transfer potential, and implicit experience. "
            "Return ONLY a valid JSON array — no preamble, no markdown fences."
        )

        items = list(jobs)
        total_batches = (len(items) + batch_size - 1) // batch_size
        print(f"[GROQ] score_jobs_with_llm — {len(items)} jobs, {total_batches} batches of {batch_size}")

        for batch_start in range(0, len(items), batch_size):
            batch_num = batch_start // batch_size + 1
            batch = items[batch_start: batch_start + batch_size]
            print(f"[GROQ] score_jobs_with_llm — batch {batch_num}/{total_batches}")

            jobs_block = ""
            for j in batch:
                jid = j['job_id']
                # Include enriched skills in the job block so Groq has full context
                skills_list = ", ".join(j.get('skills', [])[:20]) or "Not specified"
                desc = j.get('full_description', j.get('description', ''))[:600]
                jobs_block += (
                    f"\nJOB_ID: {jid}\n"
                    f"Title   : {j.get('title', '?')}\n"
                    f"Company : {j.get('company', '?')}\n"
                    f"Exp Req : {j.get('required_experience', '?')} years\n"
                    f"Skills  : {skills_list}\n"
                    f"Description:\n\"\"\"\n{desc}\n\"\"\"\n"
                )

            user = f"""{candidate_summary}

---
JOBS TO SCORE
{jobs_block}
---

For EACH job above, return a JSON array element with:
  - "job_id"  : the exact JOB_ID string (e.g. "job_0")
  - "score"   : float 0.0–1.0 (how well candidate fits this job)
  - "reason"  : one concise sentence explaining the score

Return the COMPLETE array covering all {len(batch)} jobs. Format:
[
  {{"job_id": "job_0", "score": 0.85, "reason": "Strong Python and ML skills match the JD; 1yr experience gap is manageable."}},
  {{"job_id": "job_1", "score": 0.40, "reason": "Adjacent domain but missing required Salesforce certification."}}
]
"""
            # Budget: ~400 tokens per job for output (score + reason per item)
            out_tokens = min(400 * len(batch), 3000)
            raw = self._call(system, user, max_tokens=out_tokens, temperature=0.05)
            if not raw:
                print(f"[GROQ] score_jobs_with_llm batch {batch_num} — API returned None, skipping")
                continue

            parsed = self._parse_json(raw)
            if not isinstance(parsed, list):
                # Sometimes model wraps in {"results": [...]} — try to unwrap
                if isinstance(parsed, dict):
                    for key in ('results', 'jobs', 'scores', 'data'):
                        if isinstance(parsed.get(key), list):
                            parsed = parsed[key]
                            break
            if not isinstance(parsed, list):
                print(f"[GROQ] score_jobs_with_llm batch {batch_num} — unexpected format: {type(parsed)}")
                continue

            for item in parsed:
                if not isinstance(item, dict):
                    continue
                jid = str(item.get('job_id', '')).strip()
                raw_score = item.get('score', 0)
                try:
                    score = max(0.0, min(1.0, float(raw_score)))
                except (TypeError, ValueError):
                    score = 0.0
                if jid:
                    llm_scores[jid] = score
                    reason = item.get('reason', '')
                    print(f"[GROQ]   {jid}: {score:.2f}  — {reason[:80]}")

        print(f"[GROQ] score_jobs_with_llm — scored {len(llm_scores)}/{len(items)} jobs")
        return llm_scores


# ============================================================================
# PDF / DOCX Text Utilities
# ============================================================================
def normalize_pdf_text(text):
    text = unicodedata.normalize('NFKC', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


# ============================================================================
# ResumeParser
# ============================================================================
class ResumeParser:
    SKILL_ALIASES = {
        "sklearn": "Scikit-learn", "scikit learn": "Scikit-learn",
        "nodejs": "Node.js", "node js": "Node.js",
        "reactjs": "React.js", "react js": "React.js",
        "vuejs": "Vue.js", "nextjs": "Next.js",
        "golang": "Go", "k8s": "Kubernetes",
        "tf": "TensorFlow", "postgres": "PostgreSQL",
        "amazon web services": "AWS", "google cloud platform": "GCP",
        "restful api": "REST API", "huggingface": "HuggingFace",
        "hugging face": "HuggingFace", "powerbi": "Power BI",
    }

    STRICT_SINGLE_LETTER = {'R', 'C', 'Go'}

    EDUCATION_MAP = [
        (['m.tech in ai', 'm.tech ai', 'artificial intelligence'], 'M.Tech in Artificial Intelligence'),
        (['m.tech in data', 'm.sc in data', 'msc data'], 'M.Sc in Data Science'),
        (['m.tech', 'm.e.', 'me in'], 'M.Tech in Computer Science'),
        (['mca', 'master of computer'], 'MCA (Master of Computer Applications)'),
        (['mba'], 'MBA'),
        (['bca', 'bachelor of computer'], 'BCA (Bachelor of Computer Applications)'),
        (['b.des', 'bdes', 'design'], 'B.Des in Design'),
        (['b.tech in mechanical', 'b.e. in mechanical'], 'B.Tech in Mechanical Engineering'),
        (['b.tech in electronics', 'ece', 'electronics'], 'B.Tech in Electronics and Communication'),
        (['b.e. in information', 'b.e in information'], 'B.E. in Information Technology'),
        (['b.sc', 'bsc', 'b.s.'], 'B.Sc in Computer Science'),
        (['b.com', 'bcom'], 'B.Com with Digital Marketing Certification'),
        (['b.tech', 'b.e.'], 'B.Tech in Computer Science'),
        (['ph.d', 'phd', 'doctorate'], 'M.Tech in Computer Science'),
    ]

    def extract_text_from_pdf(self, path):
        chunks = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = self._page_text(page)
                if t:
                    chunks.append(t)
        return normalize_pdf_text("\n\n".join(chunks))

    def _page_text(self, page):
        words = page.extract_words(x_tolerance=2, y_tolerance=2) or []
        if words:
            s = self._from_words(page.width, words)
            if s and len(s.strip()) >= 30:
                return s
        return page.extract_text(layout=True) or page.extract_text() or ""

    def _from_words(self, pw, words):
        cs = [(w["x0"] + w["x1"]) / 2 for w in words]
        left = sum(1 for c in cs if c < pw * 0.45)
        right = sum(1 for c in cs if c > pw * 0.55)
        mid = len(cs) - left - right
        if (left / len(cs) > 0.22 and right / len(cs) > 0.22 and mid / len(cs) < 0.25):
            sx = pw * 0.5
            return "\n".join(
                self._lines([w for w in words if w["x1"] <= sx + 6]) + [""] +
                self._lines([w for w in words if w["x0"] >= sx - 6])
            ).strip()
        return "\n".join(self._lines(words)).strip()

    def _lines(self, words, y_tol=2.5):
        if not words:
            return []
        rows = []
        for word in sorted(words, key=lambda w: (round(w["top"], 1), w["x0"])):
            placed = False
            for row in rows:
                if abs(row["top"] - word["top"]) <= y_tol:
                    row["words"].append(word)
                    row["top"] = (row["top"] + word["top"]) / 2
                    placed = True
                    break
            if not placed:
                rows.append({"top": word["top"], "words": [word]})
        return [" ".join(w["text"] for w in sorted(r["words"], key=lambda w: w["x0"])).strip()
                for r in sorted(rows, key=lambda r: r["top"])
                if " ".join(w["text"] for w in r["words"]).strip()]

    def extract_text_from_docx(self, path):
        doc = docx.Document(path)
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    def _regex_extract_skills(self, text: str) -> list[str]:
        tl = ' ' + text.lower() + ' '
        found = set()
        for skill in SKILLS_DATABASE:
            sl = skill.lower()
            if skill in self.STRICT_SINGLE_LETTER:
                pat = r'(?<![a-zA-Z0-9])' + re.escape(sl) + r'(?![a-zA-Z0-9])'
                skill_context = re.search(
                    r'(skills?|technologies|tech stack|languages?|proficient|'
                    r'expertise|tools?|frameworks?)[^\n]{0,200}' + pat, tl, re.I)
                direct = re.search(
                    r'(?:^|[\s,•|/])' + re.escape(sl) + r'(?:[\s,•|/\n]|$)', tl)
                if not (skill_context or direct):
                    continue
            elif re.search(r'[^a-z0-9]', sl):
                pat = r'(?<![a-z0-9])' + re.escape(sl) + r'(?![a-z0-9])'
            else:
                pat = r'\b' + re.escape(sl) + r'\b'
            if re.search(pat, tl):
                found.add(skill)
        for alias, canon in self.SKILL_ALIASES.items():
            if re.search(r'(?<![a-z0-9])' + re.escape(alias) + r'(?![a-z0-9])', tl):
                found.add(canon)
        return sorted(found)

    def _regex_extract_experience(self, text: str) -> float:
        tl = text.lower()
        MONTH_MAP = {
            'jan': 1, 'january': 1, 'feb': 2, 'february': 2, 'mar': 3, 'march': 3,
            'apr': 4, 'april': 4, 'may': 5, 'jun': 6, 'june': 6, 'jul': 7, 'july': 7,
            'aug': 8, 'august': 8, 'sep': 9, 'september': 9, 'oct': 10, 'october': 10,
            'nov': 11, 'november': 11, 'dec': 12, 'december': 12,
        }
        MP = (r'(january|february|march|april|may|june|july|august|september|'
              r'october|november|december|jan|feb|mar|apr|jun|jul|aug|sep|oct|nov|dec)\.?')
        EP = (r'(january|february|march|april|may|june|july|august|september|'
              r'october|november|december|jan|feb|mar|apr|jun|jul|aug|sep|oct|nov|'
              r'dec|present|current)\.?\s*(\d{4})?')
        now = datetime.now()
        ranges = []
        for line in tl.split('\n'):
            line = re.sub(r'(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)(\d{4})',
                          r'\1 \2', line.strip(), flags=re.I)
            m2 = re.search(MP + r'\s*(\d{4})\s*[-\u2013\u2014.to]+\s*' + EP, line, re.I)
            if m2:
                try:
                    sm = MONTH_MAP.get(m2.group(1).replace('.', '').strip(), 1)
                    sy = int(m2.group(2))
                    et = m2.group(3).replace('.', '').strip()
                    eys = m2.group(4)
                    if et in ('present', 'current'):
                        em, ey = now.month, now.year
                    else:
                        em = MONTH_MAP.get(et, 12)
                        ey = int(eys) if eys else now.year
                    s, e = datetime(sy, sm, 1), datetime(ey, em, 1)
                    if e >= s:
                        ranges.append((s, e))
                except Exception:
                    continue
        if not ranges:
            return 0.0
        ranges.sort(key=lambda x: x[0])
        merged = [ranges[0]]
        for cs, ce in ranges[1:]:
            ls, le = merged[-1]
            if cs <= le + timedelta(days=1):
                merged[-1] = (ls, max(le, ce))
            else:
                merged.append((cs, ce))
        return round(sum((e - s).days for s, e in merged) / 365.25, 1)

    def _regex_extract_education(self, text: str) -> str:
        t = text.lower()
        for keywords, label in self.EDUCATION_MAP:
            if any(k in t for k in keywords):
                return label
        return 'B.Tech in Computer Science'

    def detect_domains_keyword(self, skills: list[str], top_n: int = 3) -> list[tuple[str, int]]:
        sl = [s.lower() for s in skills]
        scores = {d: sum(1 for kw in kws if any(kw in s for s in sl))
                  for d, kws in DOMAIN_KEYWORDS.items()}
        ranked = sorted(scores.items(), key=lambda x: x[1], reverse=True)
        top = [(d, s) for d, s in ranked if s > 0][:top_n]
        return top if top else [('Software Engineer', 0)]

    @staticmethod
    def _normalize_llm_skills(llm_skills: list[str]) -> list[str]:
        db_lower = {s.lower(): s for s in SKILLS_DATABASE}
        result = []
        seen = set()
        for s in llm_skills:
            s = s.strip()
            if not s or s.lower() in seen:
                continue
            seen.add(s.lower())
            canonical = db_lower.get(s.lower(), s)
            result.append(canonical)
        return sorted(result)

    def parse_resume(self, path: str, ftype: str = 'pdf',
                     groq_client: GroqLLMClient | None = None) -> dict:
        if ftype == 'pdf':
            text = self.extract_text_from_pdf(path)
            raw = ""
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        raw += t + "\n"
            raw = normalize_pdf_text(raw)
        elif ftype in ('docx', 'doc'):
            text = self.extract_text_from_docx(path)
            raw = text
        else:
            text = open(path, 'r', encoding='utf-8', errors='ignore').read()
            raw = text

        if not text or len(text) < 50:
            raise ValueError("Could not extract text from resume")

        llm_success = False
        parse_source = "keyword-regex"

        # PASS 1: Groq LLM parsing (PRIMARY)
        if groq_client:
            print("[RESUME] PASS 1: Calling Groq LLM for resume parsing ...")
            llm_result = groq_client.parse_resume_llm(text)
            if llm_result:
                print("[RESUME] PASS 1: Groq LLM returned results — extracting fields ...")
                skills = self._normalize_llm_skills(llm_result['skills'])
                experience_years = llm_result['experience_years']
                education_level = llm_result['education_level'] or self._regex_extract_education(text)
                job_title = llm_result['job_title']
                domains_raw = llm_result['domains_raw']
                domains = groq_client.validate_domains(domains_raw, skills)

                if len(skills) < 5:
                    regex_skills = self._regex_extract_skills(text)
                    seen = {s.lower() for s in skills}
                    for rs in regex_skills:
                        if rs.lower() not in seen:
                            skills.append(rs)
                            seen.add(rs.lower())

                llm_success = True
                parse_source = "groq-llm"
                print(f"[RESUME] PASS 1 SUCCESS — {len(skills)} skills, domain={domains[0][0]}, "
                      f"seniority={llm_result.get('seniority', '?')}")
            else:
                print("[RESUME] PASS 1 FAILED — Groq returned None, falling back to regex")
        else:
            print("[RESUME] Groq client not available — using regex only")

        # PASS 2: Regex / keyword fallback
        if not llm_success:
            print("[RESUME] PASS 2: Using regex/keyword fallback ...")
            skills = self._regex_extract_skills(text)
            experience_years = self._regex_extract_experience(raw)
            education_level = self._regex_extract_education(text)
            domains = self.detect_domains_keyword(skills, top_n=3)
            job_title = ""
            parse_source = "keyword-regex"
            print(f"[RESUME] PASS 2 DONE — {len(skills)} skills via regex")

        domain = domains[0][0]
        candidate_location = ''
        if llm_success and llm_result.get('location'):
            candidate_location = llm_result['location']
        return {
            'skills': skills,
            'experience_years': experience_years,
            'education_level': education_level,
            'domain': domain,
            'domains': domains,
            'job_title': job_title,
            'seniority': llm_result.get('seniority', '') if llm_success else '',
            'resume_text': text,
            'parse_source': parse_source,
            'location': candidate_location,
        }

    def parse_resume_from_text(self, text: str,
                               groq_client: GroqLLMClient | None = None) -> dict:
        """Parse resume from already-extracted text (no file needed)."""
        if not text or len(text) < 50:
            raise ValueError("Resume text is too short or empty")

        llm_success = False
        parse_source = "keyword-regex"

        if groq_client:
            llm_result = groq_client.parse_resume_llm(text)
            if llm_result:
                skills = self._normalize_llm_skills(llm_result['skills'])
                experience_years = llm_result['experience_years']
                education_level = llm_result['education_level'] or self._regex_extract_education(text)
                job_title = llm_result['job_title']
                domains_raw = llm_result['domains_raw']
                domains = groq_client.validate_domains(domains_raw, skills)
                if len(skills) < 5:
                    regex_skills = self._regex_extract_skills(text)
                    seen = {s.lower() for s in skills}
                    for rs in regex_skills:
                        if rs.lower() not in seen:
                            skills.append(rs)
                            seen.add(rs.lower())
                llm_success = True
                parse_source = "groq-llm"

        if not llm_success:
            skills = self._regex_extract_skills(text)
            experience_years = self._regex_extract_experience(text)
            education_level = self._regex_extract_education(text)
            domains = self.detect_domains_keyword(skills, top_n=3)
            job_title = ""
            parse_source = "keyword-regex"

        domain = domains[0][0]
        candidate_location = ''
        if llm_success and llm_result.get('location'):
            candidate_location = llm_result['location']
        return {
            'skills': skills,
            'experience_years': experience_years,
            'education_level': education_level,
            'domain': domain,
            'domains': domains,
            'job_title': job_title,
            'resume_text': text,
            'parse_source': parse_source,
            'location': candidate_location,
        }


# ============================================================================
# JobProcessor
# ============================================================================
class JobProcessor:
    DOMAIN_SEARCH_TERMS = {
        'Data Science': 'Data Scientist', 'Web Development': 'Web Developer',
        'DevOps': 'DevOps Engineer', 'Mobile Development': 'Mobile App Developer',
        'Backend': 'Backend Developer', 'Data Engineering': 'Data Engineer',
        'Cybersecurity': 'Cybersecurity Analyst', 'UI/UX Design': 'UI UX Designer',
        'Blockchain': 'Blockchain Developer', 'Game Development': 'Game Developer',
        'Embedded Systems': 'Embedded Systems Engineer', 'Finance': 'Finance Analyst',
        'Accounting': 'Accountant', 'Marketing': 'Marketing Executive',
        'Sales': 'Sales Executive', 'HR': 'HR Executive',
        'Operations': 'Operations Manager', 'Project Management': 'Project Manager',
        'Product Management': 'Product Manager', 'Consulting': 'Management Consultant',
        'Business Analysis': 'Business Analyst', 'Legal': 'Legal Associate',
        'Healthcare': 'Healthcare Professional', 'Education': 'Teacher',
        'Engineering': 'Mechanical Engineer', 'Creative': 'Graphic Designer',
        'Hospitality': 'Hotel Management', 'Fitness & Wellness': 'Personal Trainer',
        'Media & Film': 'Video Editor', 'Fashion': 'Fashion Designer',
        'Nutrition': 'Nutritionist', 'Account Management': 'Account Manager',
        'Software Engineer': 'Software Engineer',
    }

    NON_TECH_BROAD_FALLBACK = {
        'Healthcare': ['Doctor', 'Physician', 'Medical Officer', 'Surgeon'],
        'HR': ['HR Manager', 'Human Resources'],
        'Marketing': ['Marketing Manager', 'Digital Marketer'],
        'Sales': ['Sales Manager', 'Business Development'],
        'Finance': ['Finance Manager', 'Financial Analyst'],
        'Accounting': ['Accountant', 'CA', 'Finance Executive'],
        'Legal': ['Lawyer', 'Legal Counsel', 'Advocate'],
        'Education': ['Teacher', 'Faculty', 'Lecturer', 'Trainer'],
        'Operations': ['Operations Manager', 'Operations Executive'],
        'Hospitality': ['Hotel Manager', 'Hospitality Executive'],
        'Creative': ['Graphic Designer', 'Content Creator'],
        'Fitness & Wellness': ['Fitness Trainer', 'Wellness Coach'],
        'Media & Film': ['Video Editor', 'Content Producer'],
        'Engineering': ['Mechanical Engineer', 'Civil Engineer'],
        'Consulting': ['Management Consultant', 'Business Consultant'],
        'Project Management': ['Project Manager', 'Program Manager'],
        'Product Management': ['Product Manager', 'Product Owner'],
    }

    # Indian states/cities → country "india" mapping
    INDIAN_STATES = {
        'andhra pradesh', 'arunachal pradesh', 'assam', 'bihar', 'chhattisgarh',
        'goa', 'gujarat', 'haryana', 'himachal pradesh', 'jharkhand', 'karnataka',
        'kerala', 'madhya pradesh', 'maharashtra', 'manipur', 'meghalaya', 'mizoram',
        'nagaland', 'odisha', 'punjab', 'rajasthan', 'sikkim', 'tamil nadu',
        'telangana', 'tripura', 'uttar pradesh', 'uttarakhand', 'west bengal',
        'delhi', 'delhi ncr', 'chandigarh', 'jammu and kashmir', 'ladakh',
    }
    INDIAN_CITIES = {
        'mumbai', 'pune', 'bangalore', 'bengaluru', 'hyderabad', 'chennai',
        'kolkata', 'delhi', 'noida', 'gurgaon', 'gurugram', 'ahmedabad',
        'jaipur', 'lucknow', 'kochi', 'thiruvananthapuram', 'indore',
        'bhopal', 'coimbatore', 'vizag', 'visakhapatnam', 'nagpur', 'surat',
        'chandigarh', 'bhubaneswar', 'patna', 'mysore', 'mangalore', 'vadodara',
    }

    # Valid country strings accepted by JobSpy
    VALID_COUNTRIES = {
        'argentina', 'australia', 'austria', 'bahrain', 'bangladesh', 'belgium',
        'brazil', 'canada', 'chile', 'china', 'colombia', 'costa rica', 'czech republic',
        'denmark', 'ecuador', 'egypt', 'finland', 'france', 'germany', 'greece',
        'hong kong', 'hungary', 'india', 'indonesia', 'ireland', 'israel', 'italy',
        'japan', 'kuwait', 'luxembourg', 'malaysia', 'mexico', 'morocco', 'netherlands',
        'new zealand', 'nigeria', 'norway', 'oman', 'pakistan', 'panama', 'peru',
        'philippines', 'poland', 'portugal', 'qatar', 'romania', 'saudi arabia',
        'singapore', 'south africa', 'south korea', 'spain', 'sweden', 'switzerland',
        'taiwan', 'thailand', 'turkey', 'ukraine', 'united arab emirates',
        'uk', 'usa', 'uruguay', 'venezuela', 'vietnam',
    }

    # Common location keywords → country
    COUNTRY_HINTS = {
        'usa': ['usa', 'united states', 'us'],
        'uk': ['uk', 'united kingdom', 'england', 'london', 'manchester', 'birmingham'],
        'canada': ['canada', 'toronto', 'vancouver', 'montreal', 'ottawa'],
        'australia': ['australia', 'sydney', 'melbourne', 'brisbane'],
        'germany': ['germany', 'berlin', 'munich', 'frankfurt'],
        'singapore': ['singapore'],
        'uae': ['uae', 'united arab emirates', 'dubai', 'abu dhabi'],
    }

    @classmethod
    def _detect_country(cls, location: str) -> str:
        """Detect the Indeed/Glassdoor country code from a location string."""
        if not location:
            return 'india'
        loc = location.lower().strip()

        # Direct match against valid countries
        if loc in cls.VALID_COUNTRIES:
            return loc

        # Check if any part is a known Indian state or city
        parts = [p.strip() for p in loc.replace(',', ' ').split()]
        full_parts = [p.strip() for p in loc.split(',')]
        for part in full_parts:
            part = part.strip().lower()
            if part in cls.INDIAN_STATES or part in cls.INDIAN_CITIES:
                return 'india'
        for word in parts:
            if word in cls.INDIAN_CITIES:
                return 'india'

        # Check country hints
        for country, keywords in cls.COUNTRY_HINTS.items():
            for kw in keywords:
                if kw in loc:
                    return country

        # Check if any part of location directly matches a valid country
        for part in full_parts:
            part = part.strip().lower()
            if part in cls.VALID_COUNTRIES:
                return part

        # Default fallback
        return 'india'

    def __init__(self):
        self.skills_db = SKILLS_DATABASE

    # Sites that reliably return results — get a higher budget & longer timeout
    RELIABLE_SITES = {'indeed', 'linkedin'}
    # Region-specific site availability
    SITE_REGIONS = {
        'zip_recruiter': {'usa', 'uk', 'canada'},
        'bayt': {'uae', 'united arab emirates', 'saudi arabia', 'qatar',
                 'kuwait', 'oman', 'bahrain', 'egypt', 'morocco'},
        'naukri': {'india'},
        'glassdoor': None,   # global but flaky — keep it, short timeout
        'google': None,      # global
        'indeed': None,      # global
        'linkedin': None,    # global
    }

    def _filter_sites_for_country(self, site_names, country):
        """Remove sites that don't serve the detected country."""
        filtered = []
        for site in site_names:
            regions = self.SITE_REGIONS.get(site)
            if regions is None or country in regions:
                filtered.append(site)
            else:
                print(f"[SCRAPE] Skipping {site} — not available for country '{country}'")
        return filtered

    def scrape_jobs(self, location="India", results_wanted=100,
                    site_names=None, domains=None, domain='Software Engineer',
                    llm_job_title: str = ""):
        # Scraping runs in a SUBPROCESS to avoid DLL conflicts between
        # PyTorch native libs and jobspy/tls_client threads on Windows.
        import subprocess
        from concurrent.futures import ThreadPoolExecutor, as_completed
        print(f"[SCRAPE] Starting job scraping — location={location}, sites={site_names}")

        worker_script = os.path.join(os.path.dirname(__file__), "scrape_worker.py")

        if site_names is None:
            site_names = ['indeed', 'linkedin', 'glassdoor', 'google', 'zip_recruiter', 'bayt', 'naukri']

        search_terms = []
        if llm_job_title and llm_job_title.strip():
            search_terms.append(llm_job_title.strip())

        if domains:
            for d, _ in domains:
                term = self.DOMAIN_SEARCH_TERMS.get(d, d)
                if term not in search_terms:
                    search_terms.append(term)
        else:
            term = self.DOMAIN_SEARCH_TERMS.get(domain, domain)
            if term not in search_terms:
                search_terms.append(term)

        primary_domain = domain if domain else (domains[0][0] if domains else '')
        for fallback_term in self.NON_TECH_BROAD_FALLBACK.get(primary_domain, []):
            if fallback_term not in search_terms:
                search_terms.append(fallback_term)

        country = self._detect_country(location)

        # Filter out sites that don't serve this country
        site_names = self._filter_sites_for_country(site_names, country)

        # Reliable sites get more results; flaky ones get a smaller budget
        per_term_reliable = 15
        per_term_other = 10

        print(f"[SCRAPE] Search terms: {search_terms}")
        print(f"[SCRAPE] Sites (after region filter): {site_names}")
        print(f"[SCRAPE] Reliable={per_term_reliable}/site, Other={per_term_other}/site, country={country}")
        all_records = []

        def _scrape_one(site, term):
            """Scrape a single site+term combo in a subprocess."""
            import time as _t
            is_reliable = site in self.RELIABLE_SITES
            per_term = per_term_reliable if is_reliable else per_term_other
            timeout = 120 if is_reliable else 45

            params = {
                "site_names": [site],
                "search_term": term,
                "location": location,
                "results_wanted": per_term,
                "country_indeed": country,
                "description_format": "markdown",
                "linkedin_fetch_description": True,
            }
            # Google Jobs needs google_search_term to find results
            if site == "google":
                params["google_search_term"] = f"{term} jobs in {location}"

            params_json = json.dumps(params)
            _t0 = _t.time()
            try:
                proc = subprocess.run(
                    [sys.executable, worker_script],
                    input=params_json,
                    capture_output=True, text=True, timeout=timeout,
                    cwd=os.path.dirname(__file__),
                )
                _elapsed = _t.time() - _t0
                if proc.returncode != 0:
                    print(f"[SCRAPE]   {site}/'{term}' CRASHED in {_elapsed:.1f}s — exit {proc.returncode}")
                    if proc.stderr:
                        print(f"[SCRAPE]     stderr: {proc.stderr.strip()[:200]}")
                    return []

                data = json.loads(proc.stdout)
                if not data.get("success"):
                    print(f"[SCRAPE]   {site}/'{term}' FAILED in {_elapsed:.1f}s: {data.get('error', 'unknown')[:120]}")
                    return []
                jobs = data.get("jobs", [])
                for j in jobs:
                    j.setdefault('site', site)
                print(f"[SCRAPE]   {site}/'{term}' OK in {_elapsed:.1f}s — {len(jobs)} jobs")
                return jobs
            except subprocess.TimeoutExpired:
                print(f"[SCRAPE]   {site}/'{term}' TIMED OUT ({timeout}s)")
                return []
            except Exception as e:
                print(f"[SCRAPE]   {site}/'{term}' EXCEPTION: {e}")
                return []

        # Scrape all sites in parallel (per search term) to maximise coverage
        # and minimise total wall-clock time.
        for term in search_terms:
            print(f"[SCRAPE] Scraping '{term}' across {len(site_names)} platforms (parallel) ...")
            with ThreadPoolExecutor(max_workers=len(site_names)) as executor:
                futures = {executor.submit(_scrape_one, site, term): site for site in site_names}
                for future in as_completed(futures):
                    jobs = future.result()
                    if jobs:
                        all_records.extend(jobs)

        if not all_records:
            return []

        # Deduplicate
        seen = set()
        unique = []
        for job in all_records:
            key = (job.get('title', ''), job.get('company', ''))
            if key not in seen:
                seen.add(key)
                unique.append(job)

        # Per-site breakdown so you can see diversity at a glance
        from collections import Counter
        site_counts = Counter(j.get('site', 'unknown') for j in unique)
        print(f"[SCRAPE] Total unique jobs: {len(unique)} — by site: {dict(site_counts)}")
        return unique

    def extract_skills_from_description(self, description):
        if not description:
            return []
        text_lower = ' ' + str(description).lower() + ' '
        found = [sk for sk in self.skills_db
                 if re.search(r'\b' + re.escape(sk.lower()) + r'\b', text_lower)]
        return sorted(set(found))

    def extract_required_experience(self, description):
        if not description:
            return 3
        tl = str(description).lower()
        for pat in [
            r'(\d+)\+?\s*(?:to|\-)\s*(\d+)\s*(?:years?|yrs?)',
            r'(\d+)\+?\s*(?:years?|yrs?)(?:\s+of)?\s+(?:experience|exp)',
        ]:
            m = re.search(pat, tl)
            if m:
                if len(m.groups()) > 1 and m.group(2):
                    return int((int(m.group(1)) + int(m.group(2))) / 2)
                return int(m.group(1))
        if any(t in tl for t in ['entry level', 'junior', 'fresher', 'graduate']):
            return 0
        if any(t in tl for t in ['senior', 'lead', 'principal', 'staff']):
            return 7
        return 3

    def process_jobs(self, jobs_list, is_non_tech: bool = False, groq_client=None):
        processed = []
        for i, job in enumerate(jobs_list):
            try:
                desc = job.get('description', '') or ''
                # Keep jobs even if the scraper didn't return a description —
                # we'd rather show diverse results from LinkedIn/Google/etc with
                # a thin description than lose them entirely. Fall back to title.
                if not desc:
                    desc = job.get('title', '') or ''
                if not desc:
                    continue
                skills = self.extract_skills_from_description(desc)
                req_exp = self.extract_required_experience(desc)
                raw_type = str(job.get('job_type', '') or '').strip().lower()
                if 'part' in raw_type:
                    jtype = 'Part-time'
                elif 'contract' in raw_type:
                    jtype = 'Contract'
                elif 'intern' in raw_type:
                    jtype = 'Internship'
                elif 'hybrid' in raw_type:
                    jtype = 'Hybrid'
                elif 'remote' in raw_type or 'wfh' in raw_type:
                    jtype = 'Work from Home'
                else:
                    jtype = 'Full-time'
                processed.append({
                    'job_id': f"job_{i}",
                    'title': job.get('title', 'Unknown Position'),
                    'company': job.get('company', 'Unknown Company'),
                    'location': job.get('location', 'Not specified'),
                    'job_type': jtype,
                    'description': str(desc)[:500],
                    'full_description': str(desc),
                    'skills': skills,
                    'required_experience': req_exp,
                    'job_url': job.get('job_url', '#'),
                    'date_posted': str(job.get('date_posted', 'Recently')),
                })
            except Exception as exc:
                print(f"[PROCESS] WARNING: skipped job {i} due to error: {exc}")
                continue

        # LLM skill enrichment — raise threshold so more jobs get enriched
        if groq_client is not None:
            enrich_threshold = 6 if is_non_tech else 8  # was 3/5 — more enrichment = better matching
            to_enrich = [
                {'job_id': j['job_id'], 'description': j['full_description']}
                for j in processed if len(j['skills']) < enrich_threshold
            ]
            print(f"[PROCESS] {len(to_enrich)}/{len(processed)} jobs need Groq skill enrichment (threshold={enrich_threshold})")
            if to_enrich:
                llm_skills_map = groq_client.extract_job_skills_llm(to_enrich)
                for job in processed:
                    jid = job['job_id']
                    if jid in llm_skills_map:
                        existing = set(s.lower() for s in job['skills'])
                        for sk in llm_skills_map[jid]:
                            if sk.lower() not in existing:
                                job['skills'].append(sk)
                                existing.add(sk.lower())

        return processed


# ============================================================================
# ConvDeepFM Model Architecture
# ============================================================================
class ConvDeepFM(nn.Module):
    def __init__(self, field_dims, deep_dim, embed_dim=64):
        super().__init__()
        num_fields = len(field_dims)
        kernel_sizes = [k for k in [2, 3, 4] if k <= num_fields]
        assert kernel_sizes, f"No valid CNN kernels for {num_fields} fields!"

        self.embedding = nn.Embedding(sum(field_dims), embed_dim)
        self.register_buffer(
            "offsets",
            torch.tensor(np.array((0, *np.cumsum(field_dims)[:-1])), dtype=torch.long)
        )
        self.embed_dropout = nn.Dropout(0.2)
        self.convs = nn.ModuleList([nn.Conv1d(embed_dim, 64, k) for k in kernel_sizes])
        cnn_out_dim = 64 * len(kernel_sizes)
        self.deep_bn = nn.BatchNorm1d(deep_dim)
        self.deep = nn.Sequential(
            nn.Linear(cnn_out_dim + deep_dim, 256), nn.ReLU(), nn.Dropout(0.4),
            nn.Linear(256, 128), nn.ReLU(), nn.Dropout(0.3),
            nn.Linear(128, 1)
        )
        self.fusion = nn.Linear(2, 1)

    def forward(self, x_cat, x_deep):
        x_deep = self.deep_bn(x_deep)
        x_cat = x_cat + self.offsets
        emb = self.embed_dropout(self.embedding(x_cat))
        sum_sq = torch.sum(emb, dim=1) ** 2
        sq_sum = torch.sum(emb ** 2, dim=1)
        fm_out = 0.5 * torch.sum(sum_sq - sq_sum, dim=1, keepdim=True)
        x = emb.permute(0, 2, 1)
        cnn_out = torch.cat([
            torch.max(torch.relu(conv(x)), dim=2)[0] for conv in self.convs
        ], dim=1)
        deep_out = self.deep(torch.cat([cnn_out, x_deep], dim=1))
        return self.fusion(torch.cat([fm_out, deep_out], dim=1)).squeeze(1)


# ============================================================================
# ConvDeepFM Job Recommender
# ============================================================================
class ConvDeepFMJobRecommender:
    KNOWN_JOB_TYPES = ['Contract', 'Full-time', 'Hybrid', 'Internship', 'Part-time', 'Work from Home']
    KNOWN_LOCATIONS = [
        'Ahmedabad, Gujarat', 'Bangalore, Karnataka', 'Bhubaneswar, Odisha',
        'Chandigarh', 'Chennai, Tamil Nadu', 'Coimbatore, Tamil Nadu',
        'Delhi NCR', 'Gurgaon, Haryana', 'Hyderabad, Telangana',
        'Indore, Madhya Pradesh', 'Jaipur, Rajasthan', 'Kochi, Kerala',
        'Kolkata, West Bengal', 'Mumbai, Maharashtra', 'Noida, Uttar Pradesh',
        'Pune, Maharashtra', 'Remote (India)', 'Vizag, Andhra Pradesh',
    ]

    NON_TECH_DOMAINS = {
        'HR', 'Marketing', 'Sales', 'Operations', 'Legal',
        'Healthcare', 'Education', 'Hospitality', 'Finance',
        'Accounting', 'Creative', 'Fitness & Wellness', 'Fashion',
        'Nutrition', 'Account Management', 'Consulting',
        'Business Analysis', 'Project Management', 'Product Management',
        'Media & Film', 'Engineering',
    }

    def __init__(self, model_dir="./", groq_api_key: str = ""):
        self.device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
        self.model_dir = model_dir

        self.model_path = os.path.join(model_dir, "convdeepfm_best.pth")
        self.meta_path = os.path.join(model_dir, "convdeepfm_meta.pth")
        self.encoders_path = os.path.join(model_dir, "label_encoders.pkl")

        self.user_id_map = {}
        self.job_id_map = {}
        self.next_user_id = 0
        self.next_job_id = 0

        self.model_weight = 0.4
        self.skill_weight = 0.6

        # Groq LLM client — accept explicit key or fall back to env var
        if not groq_api_key or groq_api_key in ("", "YOUR_GROQ_API_KEY_HERE"):
            groq_api_key = os.environ.get("GROQ_API_KEY", "").strip()
        if groq_api_key and groq_api_key not in ("", "YOUR_GROQ_API_KEY_HERE"):
            self.groq = GroqLLMClient(api_key=groq_api_key, model=GROQ_MODEL)
            print("Groq LLM client initialised")
        else:
            self.groq = None
            print("No Groq API key — LLM parsing disabled (keyword-regex only)")

        self.load_trained_model()

    def load_trained_model(self):
        print("Loading ConvDeepFM model & artifacts...")
        self.metadata = torch.load(self.meta_path, map_location=self.device, weights_only=False)
        bert_model_name = self.metadata.get('bert_model', 'all-MiniLM-L6-v2')
        self.bert = SentenceTransformer(bert_model_name)
        self.bert_dim = self.metadata.get('bert_dim', 384)
        self.model = ConvDeepFM(
            field_dims=self.metadata['field_dims'],
            deep_dim=self.metadata['deep_dim'],
            embed_dim=self.metadata['embed_dim']
        ).to(self.device)
        self.model.load_state_dict(
            torch.load(self.model_path, map_location=self.device, weights_only=True))
        self.model.eval()
        with open(self.encoders_path, 'rb') as f:
            self.label_encoders = pickle.load(f)
        print("All model artifacts loaded!")

    def safe_encode(self, encoder, value):
        try:
            return encoder.transform([value])[0]
        except ValueError:
            val_lower = str(value).lower()
            for cls in encoder.classes_:
                if val_lower in cls.lower() or cls.lower() in val_lower:
                    try:
                        return encoder.transform([cls])[0]
                    except ValueError:
                        pass
            city = val_lower.split(',')[0].strip()
            for cls in encoder.classes_:
                if city and city in cls.lower():
                    try:
                        return encoder.transform([cls])[0]
                    except ValueError:
                        pass
            return 0

    def _bert_encode(self, text):
        return self.bert.encode(
            [text], batch_size=1, show_progress_bar=False,
            convert_to_numpy=True, normalize_embeddings=True)[0]

    def prepare_features(self, user_profile, job):
        x_cat = np.array([
            0,  # user_id — unseen at inference
            0,  # job_id — unseen at inference
            self.safe_encode(self.label_encoders['user_education'], user_profile['education_level']),
            self.safe_encode(self.label_encoders['job_type'], job['job_type']),
            self.safe_encode(self.label_encoders['location'], job['location']),
        ], dtype=np.int64)

        user_text = (user_profile.get('resume_text', '') + " " +
                     " ".join(user_profile.get('skills', [])))
        job_text = (job.get('full_description', job['description']) + " " +
                    " ".join(job.get('skills', [])))

        user_emb = self._bert_encode(user_text)
        job_emb = self._bert_encode(job_text)

        numeric = np.array([
            user_profile['experience_years'],
            job['required_experience'],
            user_profile['experience_years'] - job['required_experience'],
        ], dtype=np.float32)

        x_deep = np.concatenate([user_emb, job_emb, numeric]).astype(np.float32)[np.newaxis, :]
        return x_cat, x_deep

    def predict_score(self, user_profile, job):
        x_cat, x_deep = self.prepare_features(user_profile, job)
        x_cat_t = torch.LongTensor(x_cat).unsqueeze(0).to(self.device)
        x_deep_t = torch.FloatTensor(x_deep).to(self.device)
        with torch.no_grad():
            raw = self.model(x_cat_t, x_deep_t).item()
        return 1 / (1 + np.exp(-raw))

    # All platforms supported by JobSpy
    ALL_SITES = ["indeed", "linkedin", "glassdoor", "google", "zip_recruiter", "bayt", "naukri"]

    def recommend_from_resume(self, resume_path, resume_type='pdf', location="India",
                              num_jobs=100, top_k=15, min_skill_match=0.05,
                              site_names=None):
        import time as _time
        pipeline_start = _time.time()
        if site_names is None:
            site_names = self.ALL_SITES

        # ── STEP 1: Parse resume ──
        print(f"\n{'='*60}")
        print(f"[PIPELINE] STEP 1/5: Parsing resume ({resume_type}) ...")
        print(f"[PIPELINE]   file: {resume_path}")
        print(f"[PIPELINE]   groq client: {'YES' if self.groq else 'NO (regex-only)'}")
        t0 = _time.time()
        parser = ResumeParser()
        user_profile = parser.parse_resume(resume_path, resume_type, groq_client=self.groq)
        print(f"[PIPELINE] STEP 1 DONE in {_time.time()-t0:.1f}s")
        print(f"[PIPELINE]   parse_source: {user_profile.get('parse_source', '?')}")
        print(f"[PIPELINE]   skills ({len(user_profile['skills'])}): {user_profile['skills'][:10]}")
        print(f"[PIPELINE]   domain: {user_profile.get('domain', '?')}")
        print(f"[PIPELINE]   job_title: {user_profile.get('job_title', '?')}")
        print(f"[PIPELINE]   experience: {user_profile.get('experience_years', '?')} years")

        # ── Auto-detect location from resume, fallback to param, fallback to India ──
        resume_location = user_profile.get('location', '').strip()
        if resume_location:
            location = resume_location
            print(f"[PIPELINE]   location (from resume): {location}")
        else:
            print(f"[PIPELINE]   location (fallback): {location}")

        if not user_profile['skills']:
            print("[PIPELINE]   WARNING: No skills detected — domain-only search")

        # ── STEP 2: Scrape jobs from ALL platforms ──
        print(f"\n[PIPELINE] STEP 2/5: Scraping jobs from {site_names} ...")
        print(f"[PIPELINE]   location: {location}, num_jobs: {num_jobs}")
        t0 = _time.time()
        processor = JobProcessor()
        raw_jobs = processor.scrape_jobs(
            location, num_jobs, site_names,
            domains=user_profile.get('domains'),
            domain=user_profile['domain'],
            llm_job_title=user_profile.get('job_title', ''),
        )
        print(f"[PIPELINE] STEP 2 DONE in {_time.time()-t0:.1f}s — {len(raw_jobs) if raw_jobs else 0} jobs scraped")

        if not raw_jobs:
            print("[PIPELINE] ABORTED — no jobs found. Try different location/domain.")
            return None

        # ── STEP 3: Process jobs (+ Groq skill enrichment) ──
        is_non_tech = user_profile.get('domain', '') in self.NON_TECH_DOMAINS
        print(f"\n[PIPELINE] STEP 3/5: Processing {len(raw_jobs)} jobs (non_tech={is_non_tech}) ...")
        print(f"[PIPELINE]   Groq skill enrichment: {'YES' if self.groq else 'NO'}")
        t0 = _time.time()
        jobs = processor.process_jobs(raw_jobs, is_non_tech=is_non_tech, groq_client=self.groq)
        print(f"[PIPELINE] STEP 3 DONE in {_time.time()-t0:.1f}s — {len(jobs) if jobs else 0} jobs processed")

        if not jobs:
            print("[PIPELINE] ABORTED — no jobs after processing.")
            return None

        # ── STEP 4: Score & rank ──
        print(f"\n[PIPELINE] STEP 4/5: Scoring & ranking {len(jobs)} jobs (top_k={top_k}) ...")
        t0 = _time.time()
        result = self._score_and_rank(user_profile, jobs, is_non_tech, top_k, location)
        print(f"[PIPELINE] STEP 4 DONE in {_time.time()-t0:.1f}s")

        # ── STEP 5: Done ──
        total = _time.time() - pipeline_start
        n_recs = len(result.get('recommendations', [])) if result else 0
        print(f"\n[PIPELINE] STEP 5/5: COMPLETE in {total:.1f}s total — {n_recs} recommendations")
        print(f"{'='*60}\n")

        return result

    def recommend_from_text(self, resume_text: str, location="India",
                            num_jobs=100, top_k=15, site_names=None):
        """Recommend jobs from already-extracted resume text."""
        if site_names is None:
            site_names = self.ALL_SITES

        parser = ResumeParser()
        user_profile = parser.parse_resume_from_text(resume_text, groq_client=self.groq)

        # Auto-detect location from resume
        resume_location = user_profile.get('location', '').strip()
        if resume_location:
            location = resume_location

        processor = JobProcessor()
        raw_jobs = processor.scrape_jobs(
            location, num_jobs, site_names,
            domains=user_profile.get('domains'),
            domain=user_profile['domain'],
            llm_job_title=user_profile.get('job_title', ''),
        )

        if not raw_jobs:
            return None

        is_non_tech = user_profile.get('domain', '') in self.NON_TECH_DOMAINS
        jobs = processor.process_jobs(raw_jobs, is_non_tech=is_non_tech, groq_client=self.groq)
        if not jobs:
            return None

        return self._score_and_rank(user_profile, jobs, is_non_tech, top_k, location)

    def _score_and_rank(self, user_profile, jobs, is_non_tech, top_k, location):
        """
        Score and rank jobs against a user profile using 4 signals:

          1. ConvDeepFM model score  — learned latent preference signal
          2. Skill overlap score     — LLM-extracted skills vs job skills
          3. BERT semantic sim       — cosine similarity of resume ↔ job embeddings
          4. Groq LLM relevance      — direct LLM judgement of candidate-job fit (NEW)

        When Groq is available, signal 4 is the strongest weight because it
        understands nuance (career trajectory, implicit experience, seniority) that
        the other three cannot. The LLM is only called on the top-2× candidates
        after a quick pre-sort using signals 1–3, to minimise API usage.
        """
        import time as _time

        user_skills_set = set(user_profile['skills'])
        user_skills_lower = set(s.lower() for s in user_skills_set)

        # ── PRE-COMPUTE user BERT embedding once ──────────────────────────────
        user_text = (user_profile.get('resume_text', '') + " " +
                     " ".join(user_profile.get('skills', [])))
        user_emb_cache = self._bert_encode(user_text)

        # ── PASS 1: Score all jobs with signals 1–3 (fast, local) ────────────
        print(f"[SCORE] PASS 1: Computing model/skill/semantic scores for {len(jobs)} jobs ...")
        t0 = _time.time()

        pre_scored = []
        for job in jobs:
            job_skills_set = set(s.lower() for s in job['skills'])

            if job_skills_set:
                skill_match = len(user_skills_lower & job_skills_set) / len(job_skills_set)
            else:
                skill_match = 0.0

            job_text = (job.get('full_description', job['description']) + " " +
                        " ".join(job.get('skills', [])))
            job_emb = self._bert_encode(job_text)
            sem_sim = max(0.0, float(np.dot(user_emb_cache, job_emb)))

            model_score = self.predict_score(user_profile, job)

            # Temporary 3-signal score for pre-sort
            if is_non_tech:
                pre_score = 0.50 * model_score + 0.15 * skill_match + 0.35 * sem_sim
            else:
                pre_score = 0.40 * model_score + 0.35 * skill_match + 0.25 * sem_sim

            pre_scored.append({
                'job': job,
                'model_score': model_score,
                'skill_match': skill_match,
                'sem_sim': sem_sim,
                'pre_score': pre_score,
            })

        pre_scored.sort(key=lambda x: x['pre_score'], reverse=True)
        print(f"[SCORE] PASS 1 done in {_time.time()-t0:.1f}s")

        # ── PASS 2: Groq LLM relevance scoring on top candidates only ─────────
        # Only score top min(top_k*3, 60) jobs to keep API calls reasonable.
        llm_score_map: dict[str, float] = {}
        llm_available = self.groq is not None

        if llm_available:
            llm_candidates_n = min(top_k * 3, 60)
            llm_candidates = [r['job'] for r in pre_scored[:llm_candidates_n]]
            print(f"[SCORE] PASS 2: Groq LLM scoring top {len(llm_candidates)} candidates ...")
            t0 = _time.time()
            llm_score_map = self.groq.score_jobs_with_llm(user_profile, llm_candidates)
            print(f"[SCORE] PASS 2 done in {_time.time()-t0:.1f}s — "
                  f"{len(llm_score_map)}/{len(llm_candidates)} jobs scored by Groq")
        else:
            print("[SCORE] PASS 2 skipped — no Groq client")

        # ── PASS 3: Final combined score + experience modifier ────────────────
        # Weights when Groq LLM score is available (4 signals):
        #   Non-tech: model=0.25, skill=0.10, sem=0.20, llm=0.45
        #   Tech    : model=0.20, skill=0.25, sem=0.10, llm=0.45
        # Weights when Groq is unavailable (3 signals, preserve original balance):
        #   Non-tech: model=0.50, skill=0.15, sem=0.35
        #   Tech    : model=0.40, skill=0.35, sem=0.25
        if llm_available and llm_score_map:
            if is_non_tech:
                mw, sw, semw, llmw = 0.25, 0.10, 0.20, 0.45
            else:
                mw, sw, semw, llmw = 0.20, 0.25, 0.10, 0.45
        else:
            llmw = 0.0
            if is_non_tech:
                mw, sw, semw = 0.50, 0.15, 0.35
            else:
                mw, sw, semw = 0.40, 0.35, 0.25

        print(f"[SCORE] Weights — model={mw:.0%} skill={sw:.0%} sem={semw:.0%} llm={llmw:.0%}")

        recommendations = []
        for entry in pre_scored:
            job = entry['job']
            model_score = entry['model_score']
            skill_match = entry['skill_match']
            sem_sim = entry['sem_sim']
            jid = job['job_id']

            llm_score = llm_score_map.get(jid, 0.0)

            final_score = (mw * model_score
                           + sw * skill_match
                           + semw * sem_sim
                           + llmw * llm_score)

            # Experience gap modifier
            req_exp = job.get('required_experience')
            if req_exp is not None and req_exp > 0:
                gap = user_profile['experience_years'] - req_exp
                if gap >= 0:
                    exp_mod = 1.15 if gap <= 2 else (1.05 if gap <= 5 else 1.0)
                else:
                    exp_mod = max(0.4, 1.0 + gap * 0.08)
            else:
                exp_mod = 1.1
            final_score = min(final_score * exp_mod, 1.0)

            recommendations.append({
                'job': job,
                'model_score': model_score,
                'skill_match': skill_match,
                'sem_sim': sem_sim,
                'llm_score': llm_score,
                'final_score': final_score,
                'matching_skills': [s for s in job['skills'] if s.lower() in user_skills_lower],
                'missing_skills':  [s for s in job['skills'] if s.lower() not in user_skills_lower],
            })

        recommendations.sort(key=lambda x: x['final_score'], reverse=True)
        recommendations = recommendations[:top_k]

        # Log final distribution
        if recommendations:
            fs = [r['final_score'] for r in recommendations]
            print(f"[SCORE] Final scores — min={min(fs):.2f} max={max(fs):.2f} "
                  f"avg={sum(fs)/len(fs):.2f}")

        result = {
            'user_profile': user_profile,
            'recommendations': recommendations,
            'total_jobs_analyzed': len(jobs),
            'location': location,
            'timestamp': datetime.now().isoformat(),
            'scoring_config': {
                'model_weight': mw,
                'skill_weight': sw,
                'semantic_weight': semw,
                'llm_weight': llmw,
                'groq_scored_jobs': len(llm_score_map),
            },
        }
        self._display_results(user_profile, recommendations, location)
        return result

    def _display_results(self, user: dict, recs: list, location: str) -> None:
        """Print a formatted summary of the recommendation results to stdout."""
        sep = "=" * 80
        print(f"\n{sep}")
        print("TOP JOB RECOMMENDATIONS")
        print(sep)
        print(f"\nCANDIDATE PROFILE:")
        print(f"  Skills     : {len(user['skills'])}")
        print(f"  Experience : {user['experience_years']} years")
        print(f"  Education  : {user['education_level']}")
        print(f"  Domain     : {user['domain']}")
        print(f"  Seniority  : {user.get('seniority', 'N/A')}")
        if user.get('job_title'):
            print(f"  Job Title  : {user['job_title']}")
        print(f"  Location   : {location}")
        print(f"  Parsed via : {user.get('parse_source', 'unknown')}")

        if recs:
            model_scores = [r['model_score'] for r in recs]
            skill_scores = [r['skill_match'] for r in recs]
            final_scores = [r['final_score'] for r in recs]
            llm_scores   = [r.get('llm_score', 0) for r in recs]
            print(f"\nSCORE DISTRIBUTION (top {len(recs)}):")
            print(f"  Model : {min(model_scores):.2f}–{max(model_scores):.2f}  "
                  f"(avg {sum(model_scores)/len(model_scores):.2f})")
            print(f"  Skill : {min(skill_scores):.2f}–{max(skill_scores):.2f}  "
                  f"(avg {sum(skill_scores)/len(skill_scores):.2f})")
            if any(s > 0 for s in llm_scores):
                print(f"  LLM   : {min(llm_scores):.2f}–{max(llm_scores):.2f}  "
                      f"(avg {sum(llm_scores)/len(llm_scores):.2f})")
            print(f"  Final : {min(final_scores):.2f}–{max(final_scores):.2f}  "
                  f"(avg {sum(final_scores)/len(final_scores):.2f})")

        print(f"\n{sep}")
        for i, rec in enumerate(recs, 1):
            job = rec['job']
            exp_req = job.get('required_experience')
            if exp_req is not None:
                gap = user['experience_years'] - exp_req
                exp_flag = "OK" if abs(gap) <= 2 else ("OVER" if gap > 0 else "UNDER")
                exp_str = f"{exp_req}yr required [{exp_flag}: you have {user['experience_years']}yr]"
            else:
                exp_str = "Experience not specified"

            llm_s = rec.get('llm_score', 0)
            llm_str = f"  LLM:{llm_s:.0%}" if llm_s > 0 else ""
            print(f"\n{i:>2}. {job['title']}")
            print(f"    {job['company']}  |  {job['location']}  |  {job['job_type']}")
            print(f"    {exp_str}")
            print(f"    Posted: {job['date_posted']}")
            print(f"    Score: {rec['final_score']:.0%}  "
                  f"(Model:{rec['model_score']:.0%}  "
                  f"Skill:{rec['skill_match']:.0%}  "
                  f"Sem:{rec.get('sem_sim', 0):.0%}{llm_str})")
            if rec['matching_skills']:
                ms = ", ".join(rec['matching_skills'][:6])
                print(f"    Matching skills ({len(rec['matching_skills'])}): {ms}")
            if rec['missing_skills'][:3]:
                gap_s = ", ".join(rec['missing_skills'][:3])
                print(f"    Skills to learn: {gap_s}")
            url = job.get('job_url', '#')
            print(f"    {url[:90]}")
            print(f"    {'-'*78}")
        print(sep)
